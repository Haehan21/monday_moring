import feedparser
from datetime import datetime, timedelta
from docx import Document
from deep_translator import GoogleTranslator
import docx  # docx 모듈 임포트

# 하이퍼링크 추가 함수
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    r_pr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


raise

# 오늘 날짜와 9일 전 날짜 계산
today = datetime.now()
start_date = today - timedelta(days=9)
today_str = today.strftime('%Y%m%d')  # 오늘 날짜를 문자열로 변환

# 검색할 키워드 리스트
keywords = {
    "미국": [
        "Logistics industry trends", "Logistics technology innovation", "Logistics market forecast",
        "Logistics automation", "Logistics startup", "Supply chain management", "Logistics cost reduction",
        "Logistics trends", "Logistics innovation cases", "Logistics blockchain", "3PL", "4PL", "Thrid Party Logistics",
        "cross border e-commerce"
    ],
    "일본": [
        "物流業界の動向", "物流技術革新", "物流市場展望", "物流自動化", "物流スタートアップ", "サプライチェーン管理", 
        "サプライチェーンかんり", "物流コスト削減", "物流トレンド", "物流革新事例", "物流ブロックチェーン", "3PL", "4PL", 
        "Thrid Party Logistics", "cross border e-commerce"
    ],
    "중국": [
        "物流行业趋势", "物流技术创新", "物流市场看法", "物流自动化", "物流初创公司", "供应链管理", "物流成本降低",
        "物流趋势", "物流创新案例", "物流区块链", "3PL", "4PL", "Thrid Party Logistics", "cross border e-commerce"
    ],
    "대한민국": [
        "물류 산업 동향", "물류 기술 혁신", "물류 시장 전망", "물류 자동화", "물류 스타트업", "물류 공급망 관리",
        "물류 트렌드", "물류 혁신 사례", "물류 블록체인", "물류 디지털 전환", "3PL", "4PL", "Thrid Party Logistics", "3자 물류",
        "cross border e-commerce"
    ]
}

# 검색할 지역 리스트
regions = [
    {"name": "미국", "hl": "en", "gl": "US"},
    {"name": "일본", "hl": "ja", "gl": "JP"},
    {"name": "중국", "hl": "zh-CN", "gl": "CN"},
    {"name": "대한민국", "hl": "ko", "gl": "KR"}
]

# 중복 제거를 위한 집합
seen_links = set()

# 번역기 설정
translator = GoogleTranslator(source='auto', target='ko')

# 각 지역에 대해 별도의 Word 문서 생성
for region in regions:
    doc = Document()
    doc.add_heading(f'{region["name"]} 뉴스 검색 결과', 0)
    
    region_keywords = keywords[region["name"]]
    for keyword in region_keywords:
        encoded_keyword = keyword.replace(" ", "%20")
        rss_url = f"https://news.google.com/rss/search?q={encoded_keyword}&hl={region['hl']}&gl={region['gl']}&ceid={region['gl']}:{region['hl']}"
        feed = feedparser.parse(rss_url)
        
        doc.add_heading(f'키워드: {keyword}', level=1)
        for entry in feed.entries:
            published_date = datetime(*entry.published_parsed[:6])
            if start_date <= published_date <= today:
                if entry.link not in seen_links and "유료" not in entry.title and "보고서" not in entry.title and "subscription" not in entry.title and "premium" not in entry.title and "Get a Sample PDF of report" not in entry.title and "샘플 보고서 다운로드" not in entry.title and "有料" not in entry.title and "报告" not in entry.title and "サンプルPDFを取得" not in entry.title:
                    seen_links.add(entry.link)
                    translated_title = translator.translate(entry.title)
                    paragraph = doc.add_paragraph()
                    add_hyperlink(paragraph, translated_title, entry.link)
                    paragraph.add_run(f", {published_date.strftime('%Y-%m-%d')}")
                    # doc.add_paragraph("-" * 50)
    
    # 각 지역별로 Word 문서 저장
    doc.save(f'{today_str}_logistics_news_{region["name"]}.docx')
    print(f"{region['name']} 뉴스 검색 결과가 '{today_str}_logistics_news_{region["name"]}.docx' 파일에 저장되었습니다.")

print("모든 뉴스 검색 결과가 저장되었습니다.")